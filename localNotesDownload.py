import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Input syllabus as a dictionary of units and topics
syllabus = {
    "Unit – 1": "AVL Tree, Binary Search Tree"
}

def preprocess_gfg_content(raw_content):
    """
    Clean and preprocess the raw content fetched from GFG.
    Args:
        raw_content (str): The raw content string from the webpage.
    Returns:
        str: Cleaned and meaningful content.
    """
    # Remove multiple newlines and excessive whitespace
    cleaned_content = re.sub(r'\n+', '\n', raw_content)  # Collapse multiple newlines
    cleaned_content = cleaned_content.strip()  # Strip leading and trailing whitespace

    # Remove phrases like "Summarize", "Comments", "Save", "Share", etc.
    to_remove = [
        r'\bSummarize\b', r'\bComments\b', r'\bImprove\b', r'\bLike Article\b', r'\bSave\b',
        r'\bShare\b', r'\bReport\b', r'\bFollow\b', r'Last Updated\s*:\s*\d{1,2} \w+, \d{4}'
    ]
    for phrase in to_remove:
        cleaned_content = re.sub(phrase, '', cleaned_content, flags=re.IGNORECASE)

    # Remove additional patterns (if needed, adjust based on results)
    cleaned_content = re.sub(r'\bSuggest changes\b', '', cleaned_content, flags=re.IGNORECASE)

    # Remove any lingering extra whitespace
    cleaned_content = re.sub(r'\s{2,}', ' ', cleaned_content)  # Replace multiple spaces with a single space
    return cleaned_content


def search_gfg_with_google(query):
    """Search GeeksforGeeks articles using their internal API"""
    search_url = f"https://recommendations.geeksforgeeks.org/api/v1/global-search?products=articles&query={query.replace(' ', '+')}&articles_count=1"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    response = requests.get(search_url, headers=headers)
    data = response.json()
    return data['detail']['articles']['data'][0]['post_url']


def fetch_gfg_article_html(url):
    """
    Fetch and parse the HTML content of a GFG article.
    Args:
        url (str): URL of the GFG article.
    Returns:
        str: Cleaned HTML content of the article.
    """
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    article = soup.find('article', {'class': 'content'})
    if not article:
        return None

    # Remove unwanted elements (e.g., ads, share buttons)
    for element in article.find_all(['script', 'style', 'nav', 'footer', 'aside', 'form']):
        element.decompose()

    # Clean and return the HTML content
    return str(article)


def download_image(image_url, save_folder):
    """
    Downloads an image and returns the local file path.
    """
    try:
        response = requests.get(image_url, stream=True)
        if response.status_code == 200:
            # Ensure the save folder exists
            os.makedirs(save_folder, exist_ok=True)

            filename = os.path.join(save_folder, image_url.split("/")[-1])
            
            # Save the image locally
            with open(filename, "wb") as file:
                for chunk in response.iter_content(1024):
                    file.write(chunk)
            
            return filename
    except Exception as e:
        print(f"Failed to download image {image_url}: {e}")
    return None

def add_images_to_doc(images, doc, save_folder):
    """
    Downloads images first and then adds them to a Word document.
    """
    for img in images:
        img_url = img.get("src")
        if img_url and img_url.startswith("http"):  # Only process valid image URLs
            local_path = download_image(img_url, save_folder)
            if local_path:
                doc.add_picture(local_path, width=Inches(5))
                doc.add_paragraph("\n")  # Space after the image

def convert_html_to_docx(html_content, output_path, save_folder="images"):
    """
    Convert HTML content to a Word document.
    
    Args:
        html_content (str): The HTML content to be converted.
        output_path (str): The path to save the Word document.
        save_folder (str): Folder where images will be saved.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    # Process each element in the HTML
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'img','pre']):
        if element.name.startswith('h'):  # Headings
            level = int(element.name[1])  # Extract heading level (1, 2, 3)
            doc.add_heading(element.get_text().strip(), level=level)
        elif element.name == 'p':  # Paragraphs
            doc.add_paragraph(element.get_text().strip())
        elif element.name in ['ul', 'ol']:  # Lists
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text().strip(), style='List Bullet' if element.name == 'ul' else 'List Number')
        elif element.name == 'img':  # Images
            img_url = element.get('src')
            if img_url  and img_url.split(".")[-1] in ["png","jpg","jpeg"]:
                img_response = requests.get(img_url)
                img_bytes = BytesIO(img_response.content)
                doc.add_picture(img_bytes, width=Inches(4.5))
                
        elif element.name == 'pre':  # Code Block
            code_text = element.get_text().strip()
            para = doc.add_paragraph(code_text)
            run = para.runs[0]
            run.font.name = "Lucida Console"  # Set font to Lucida Console
            para.alignment = 0  # Left align
            para.style = doc.styles["Normal"]

            # Add a gray background to the code block
            shading_elm = parse_xml(r'<w:shd {} w:fill="EAEAEA"/>'.format(nsdecls('w')))
            para._element.get_or_add_pPr().append(shading_elm)

    # Save the document
    doc.save(output_path)
    print(f"Document saved at: {output_path}")
    
    
def main():
    """
    Main function to process the syllabus, fetch articles, and create Word documents.
    """
    # Create an output folder for the documents
    output_folder = "Syllabus_Notes"
    os.makedirs(output_folder, exist_ok=True)

    for unit, topics in syllabus.items():
        unit_doc_path = os.path.join(output_folder, f"{unit.replace(' ', '_')}.docx")
        print(f"Processing {unit}...")

        # Combine all topics into a single Word document for the unit
        combined_html = f"<h1>{unit}</h1>"
        topics_list = topics.split(", ")
        for topic in topics_list:
            print(f"Searching for topic: {topic}")
            gfg_url = search_gfg_with_google(topic)
            if not gfg_url:
                print(f"No relevant GFG article found for topic: {topic}")
                combined_html += f"<h2>{topic}</h2><p>No content found.</p>"
                continue

            print(f"Found URL: {gfg_url}")
            article_html = fetch_gfg_article_html(gfg_url)

            if article_html:
                combined_html += f"<h2>{topic}</h2>{article_html}"
            else:
                print(f"No content found in article for topic: {topic}")
                combined_html += f"<h2>{topic}</h2><p>No content found in article.</p>"

        # Convert combined HTML to a Word document
        print(f"Creating document for {unit}...")
        convert_html_to_docx(combined_html, unit_doc_path)
        print(f"Document for {unit} saved at {unit_doc_path}")

    print("All units processed. Check the Syllabus_Notes folder for documents.")


if __name__ == "__main__":
    main()