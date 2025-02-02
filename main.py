import os
import re
import requests
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import streamlit as st
from agno.models.groq import Groq

import os
from dotenv import load_dotenv
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

# ----- Syllabus Notes Utilities -----

# Preprocessing text (removes extra spaces, unwanted phrases, etc.)
def preprocess_gfg_content(raw_content):
    cleaned_content = re.sub(r'\n+', '\n', raw_content)  # Collapse multiple newlines
    cleaned_content = cleaned_content.strip()  # Strip leading and trailing whitespace
    to_remove = [
        r'\bSummarize\b', r'\bComments\b', r'\bImprove\b', r'\bLike Article\b', r'\bSave\b',
        r'\bShare\b', r'\bReport\b', r'\bFollow\b', r'Last Updated\s*:\s*\d{1,2} \w+, \d{4}'
    ]
    for phrase in to_remove:
        cleaned_content = re.sub(phrase, '', cleaned_content, flags=re.IGNORECASE)
    cleaned_content = re.sub(r'\s{2,}', ' ', cleaned_content)
    return cleaned_content

# Use GeeksforGeeks internal API to search for an article URL
def search_gfg_with_google(query):
    search_url = f"https://recommendations.geeksforgeeks.org/api/v1/global-search?products=articles&query={query.replace(' ', '+')}&articles_count=1"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    response = requests.get(search_url, headers=headers)
    data = response.json()
    try:
        return data['detail']['articles']['data'][0]['post_url']
    except (KeyError, IndexError):
        return None

# Fetch article HTML from a given URL and remove unwanted elements
def fetch_gfg_article_html(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    article = soup.find('article', {'class': 'content'})
    if not article:
        return None
    for element in article.find_all(['script', 'style', 'nav', 'footer', 'aside', 'form']):
        element.decompose()
    return str(article)


def convert_html_to_docx_bytes(html_content):
    """
    Convert HTML content to a DOCX file in memory (BytesIO object).
    Processes headings, paragraphs, lists, images, and <pre> code blocks.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'img', 'pre']):
        if element.name in ['h1', 'h2', 'h3']:
            level = int(element.name[1])
            doc.add_heading(element.get_text().strip(), level=level)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text().strip())
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                style = 'List Bullet' if element.name == 'ul' else 'List Number'
                doc.add_paragraph(li.get_text().strip(), style=style)
        elif element.name == 'img':
            img_url = element.get('src')
            if img_url  and img_url.split(".")[-1] in ["png","jpg","jpeg"]:
                img_response = requests.get(img_url)
                img_bytes = BytesIO(img_response.content)
                doc.add_picture(img_bytes, width=Inches(4.5))
                doc.add_paragraph("\n")
        elif element.name == 'pre':
            # Add code block with Lucida Console font and a light gray background.
            code_text = element.get_text().strip()
            para = doc.add_paragraph(code_text)
            if para.runs:
                run = para.runs[0]
                run.font.name = "Lucida Console"
            shading_elm = parse_xml(r'<w:shd {} w:fill="EAEAEA"/>'.format(nsdecls('w')))
            para._element.get_or_add_pPr().append(shading_elm)

    # Save the DOCX document to a BytesIO stream.
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ----- Streamlit Interface -----

# Configure the Streamlit page
st.set_page_config(page_title="Study Notes & Agents", layout="wide")

# Use Streamlit tabs (or radio buttons) to switch between two functions:
tab1, tab2 = st.tabs(["Syllabus Notes Generator", "Study Agents"])

# ------- Tab 1: Syllabus Notes Generator -------
with tab1:
    st.header("Syllabus Notes Generator")
    st.write("Add units and topics. The app will fetch content from GeeksforGeeks and create a DOCX file with the notes.")
    
    # Input: Allow user to add multiple units and topics (comma separated topics per unit)
    units_input = st.text_area(
        "Enter syllabus units and topics (one unit per line, format: UnitName: topic1, topic2, ...)",
        "Unit – 1: AVL Tree, Binary Search Tree"
    )
    
    if st.button("Generate Syllabus Notes"):
        # Process each line as a unit
        syllabus = {}
        for line in units_input.splitlines():
            if ":" in line:
                unit, topics = line.split(":", 1)
                syllabus[unit.strip()] = topics.strip()
        if not syllabus:
            st.error("Please enter valid units and topics.")
        else:
            
            # Build combined HTML for each unit
            combined_docx = BytesIO()  # We'll combine the DOCX files in memory if needed.
            for unit, topics in syllabus.items():
                combined_html = f"<h1>{unit}</h1>"
                topics_list = [t.strip() for t in topics.split(",")]
                for topic in topics_list:
                    st.write(f"Searching for topic: **{topic}**...")
                    gfg_url = search_gfg_with_google(topic)
                    if not gfg_url:
                        combined_html += f"<h2>{topic}</h2><p>No content found.</p>"
                    else:
                        st.write(f"Found URL: {gfg_url}")
                        article_html = fetch_gfg_article_html(gfg_url)
                        if article_html:
                            # Optionally preprocess article content here
                            cleaned_html = preprocess_gfg_content(article_html)
                            combined_html += f"<h2>{topic}</h2>{cleaned_html}"
                        else:
                            combined_html += f"<h2>{topic}</h2><p>No content found in article.</p>"
                

                # Convert the combined HTML to a DOCX file in memory.
                docx_bytes = convert_html_to_docx_bytes(combined_html)
                
                
                
                if(docx_bytes):
                    st.success("Notes created successfully")
                # Offer a download button for this unit.
                st.download_button(
                    label=f"Download {unit} DOCX",
                    data=docx_bytes,
                    file_name=f"{unit.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            

# ------- Tab 2: Study Agents -------
with tab2:
    # Groq API Key Input
    groq_api_key = st.sidebar.text_input("Enter Your Groq API Key", type="password", key="groq_api_key_tab2")
    st.sidebar.markdown(
        "[Get your Groq API Key from the Groq Console](https://console.groq.com/keys)"
    )
    
    if groq_api_key:
        try:
            from agno.agent import Agent
            from agno.tools.youtube import YouTubeTools
            from agno.tools.arxiv import ArxivTools
            # from agno.tools.tavily import TavilyTools
            from agno.tools.newspaper4k import Newspaper4kTools

            # Create tabs for each agent
            tab_yt, tab_arxiv, tab_web, tab_flashcards = st.tabs([
                "YouTube Summarizer", 
                "Research Paper Summarizer", 
                "Web Content Summarizer", 
                "Flashcard Generator"
            ])

            # YouTube Summarizer Tab
            with tab_yt:
                st.subheader("YouTube Video Summarizer")
                video_url = st.text_input("Enter YouTube video URL", "https://www.youtube.com/watch?v=Iv9dewmcFbs&t", key="youtube_url")
                youtube_query = st.text_area("Enter your query for the video", "Summarize this video in 5 bullet points.", key="youtube_query")
                if st.button("Run YouTube Agent", key="youtube_button"):
                    with st.spinner("Processing YouTube video..."):
                        youtube_agent = Agent(
                            model=Groq(id="llama-3.3-70b-versatile", api_key=groq_api_key),
                            tools=[YouTubeTools()],
                            show_tool_calls=True,
                            description="You are a YouTube agent. Obtain the captions of a YouTube video and answer questions.",
                        )
                        response = youtube_agent.run(f"{youtube_query} {video_url}", markdown=True)
                        if response:
                            st.markdown("### YouTube Video Summary:")
                            st.markdown(response.content)
                        else:
                            st.error("No response from the YouTube agent.")

            # Research Paper Summarizer Tab
            with tab_arxiv:
                st.subheader("Research Paper Summarizer")
                arxiv_query = st.text_input("Enter Arxiv search query (e.g., 'machine learning')", "machine learning", key="arxiv_query")
                if st.button("Run Arxiv Agent", key="arxiv_button"):
                    with st.spinner("Fetching and summarizing research papers..."):
                        arxiv_agent = Agent(
                            model=Groq(id="llama-3.3-70b-versatile", api_key=groq_api_key),
                            tools=[ArxivTools()],
                            show_tool_calls=True,
                            description="You are an Arxiv agent. Fetch and summarize research papers.",
                        )
                        response = arxiv_agent.run(f"Find and summarize the 5 latest papers on {arxiv_query}.", markdown=True)
                        if response:
                            st.markdown("### Research Paper Summaries:")
                            st.markdown(response.content)
                        else:
                            st.error("No response from the Arxiv agent.")

            # Web Content Summarizer Tab
            with tab_web:
                st.subheader("Web Content Summarizer")
                web_query = st.text_input("Enter web search query (e.g., 'latest advancements in AI')", "latest advancements in AI", key="web_query")
                if st.button("Run Web Agent", key="web_button"):
                    with st.spinner("Searching and summarizing web content..."):
                        web_agent = Agent(
                            model=Groq(id="llama-3.3-70b-versatile", api_key=groq_api_key),
                            tools=[Newspaper4kTools()],
                            show_tool_calls=True,
                            description="You are a web agent. Search and summarize web content.",
                        )
                        response = web_agent.run(f"Search and summarize content about {web_query}.", markdown=True)
                        if response:
                            st.markdown("### Web Content Summary:")
                            st.markdown(response.content)
                        else:
                            st.error("No response from the web agent.")

            # Flashcard Generator Tab
            with tab_flashcards:
                st.subheader("Flashcard Generator")
                flashcard_topic = st.text_input("Enter a topic to generate flashcards", "machine learning", key="flashcard_topic")
                if st.button("Generate Flashcards", key="flashcard_button"):
                    with st.spinner("Generating flashcards..."):
                        flashcard_agent = Agent(
                            model=Groq(id="llama-3.3-70b-versatile", api_key=groq_api_key),
                            tools=[],
                            show_tool_calls=True,
                            description="You are a flashcard generator. Create flashcards for the given topic.",
                        )
                        response = flashcard_agent.run(f"Generate 5 flashcards for the topic: {flashcard_topic}.", markdown=True)
                        if response:
                            st.markdown("### Flashcards:")
                            st.markdown(response.content)
                        else:
                            st.error("No response from the flashcard agent.")

        except Exception as e:
            st.error(f"Error loading agents: {e}")
            
    else:
        st.warning("Please Add Groq Api Key")
        
footer="""<style>
a:link , a:visited{

color: blue;
background-color: transparent;
text-decoration: none;
font-weight:bold;
letter-spacing:2px;
}

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 0;
bottom: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
}

</style>
<div class="footer">
<p>Developed with ❤ by <a style='display: block; text-align: center;' href="https://linkedin.com/in/swikrit-shukla" target="_blank">Swikrit Shukla</a></p>
</div>
"""
st.markdown(footer,unsafe_allow_html=True)
